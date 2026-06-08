#!/usr/bin/env python
"""Combine the Business + Technical Markdown docs into one Word (.docx) file.

Order: Business Documentation first, then Technical Documentation.
Renders headings, tables, fenced code blocks (monospace, for ASCII diagrams),
bullet/numbered lists, blockquotes, and inline **bold** / *italic* / `code`.

Usage:
    python scripts/build_combined_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BUSINESS = ROOT / "docs" / "BUSINESS_DOCUMENTATION.md"
TECHNICAL = ROOT / "docs" / "TECHNICAL_DOCUMENTATION.md"
OUT = ROOT / "docs" / "CogLoad_Documentation.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)")


def add_runs(paragraph, text: str) -> None:
    """Apply inline **bold** / *italic* / `code` formatting to a paragraph."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Courier New"
        else:  # *italic*
            paragraph.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc, lines: list[str]) -> None:
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8)


def add_table(doc, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            para = cells[ci].paragraphs[0]
            add_runs(para, text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep_row(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def render_markdown(doc, md: str) -> None:
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, buf)
            continue

        # table: current line has pipe and next line is a separator
        if "|" in line and i + 1 < n and is_sep_row(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2  # skip header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # headings
        if stripped.startswith("#"):
            m = re.match(r"(#{1,6})\s+(.*)", stripped)
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip("> ").rstrip())
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # bullet list (one nesting level)
        mb = re.match(r"(\s*)[-*]\s+(.*)", line)
        if mb:
            indent = len(mb.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs(p, mb.group(2))
            i += 1
            continue

        # numbered list — keep literal number for reliable ordering
        mn = re.match(r"(\s*)(\d+)\.\s+(.*)", line)
        if mn:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run(f"{mn.group(2)}. ").bold = True
            add_runs(p, mn.group(3))
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # paragraph — gather wrapped continuation lines
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"\s*(#|```|>|[-*]\s|\d+\.\s)", lines[i]
        ) and "|" not in lines[i]:
            buf.append(lines[i])
            i += 1
        para = doc.add_paragraph()
        add_runs(para, " ".join(s.strip() for s in buf))


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CogLoad — Cognitive Load Detection")
    r.bold = True
    r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Combined Documentation — Business & Technical")
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_page_break()

    render_markdown(doc, BUSINESS.read_text())
    doc.add_page_break()
    render_markdown(doc, TECHNICAL.read_text())

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
